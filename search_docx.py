import os
import glob
import docx

# flatten the list for the documents
def flatten_list(l):
    return [item for sublist in l for item in sublist]

# uses glob to join path with all files ending with the .docx extention
def get_docx_files():
    # get current working directory
    path = os.getcwd()
    docx_files = []
    for folder,_,_ in os.walk(path):
        docx_files.append((glob.glob(os.path.join(folder, "*.docx"))))
    return flatten_list(docx_files)

def search_through_document(str, doc, file_name):
    # A document has paragraphs, so we iterate through every paragraph and track the paragraph number
    # That way we can let the user know where exactly the string in question is
    for paragraph_number in range(len((doc.paragraphs))):
        if str in doc.paragraphs[paragraph_number].text.lower():
            print(f'File: {file_name}; paragraph number {paragraph_number+1}')

def main():
    docx_files = get_docx_files()
    while True:
        input_string = input('What are you looking for: ').strip().lower()
        if input_string == 'end of work':
            break
        for docx_file in docx_files:
            # file_name will take the last 2 parameters of the path, which are the parent folder and the file name
            file_name = docx_file.split("\\")[-2:]
            try:
                document = docx.Document(docx_file)
                search_through_document(input_string, document, file_name)
            # handle potential errors while opening documents
            # the choices for errors to handle came from experience working with excel files
            except PermissionError:
                print(f'NO PERMISSION TO OPEN {file_name}')
            except ValueError:
                print(f'FAILED TO OPEN {file_name}')
            except:
                # the docx library has custom errors, so the final except is crucial for this library
                print(f'UNKNOWN ERROR FOR {file_name}')
            


if __name__ == "__main__":
    main()
